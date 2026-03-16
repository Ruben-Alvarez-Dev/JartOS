"""
DocumentParser - Extract text and metadata from PDF and DOCX files.
"""

import re
import logging
from pathlib import Path
from typing import Optional, List, Dict, Any, Tuple
from dataclasses import dataclass, field

logger = logging.getLogger(__name__)


@dataclass
class ParsedPage:
    """Represents a parsed page from a document."""

    page_number: int
    text: str
    tables: List[List[List[str]]] = field(default_factory=list)
    metadata: dict = field(default_factory=dict)


@dataclass
class ParsedDocument:
    """Represents a fully parsed document."""

    filename: str
    file_type: str
    title: Optional[str] = None
    tema: Optional[int] = None
    total_pages: int = 0
    pages: List[ParsedPage] = field(default_factory=list)
    full_text: str = ""
    metadata: dict = field(default_factory=dict)


class DocumentParser:
    """Extract text and metadata from PDF and DOCX documents."""

    # Metadata extraction patterns
    METADATA_PATTERNS = {
        "tema": r"(?i)tema\s*(\d+)",
        "titulo": r"(?i)t[ií]tulo[:\s]+([^\n]+)",
        "apartado": r"(?i)apartado\s*(\d+)",
    }

    def __init__(
        self,
        extract_tables: bool = True,
        extract_images: bool = False,
        metadata_patterns: Optional[Dict[str, str]] = None,
    ):
        """
        Initialize the parser.

        Args:
            extract_tables: Whether to extract tables from PDFs
            extract_images: Whether to extract images (not implemented)
            metadata_patterns: Custom regex patterns for metadata extraction
        """
        self.extract_tables = extract_tables
        self.extract_images = extract_images
        self.metadata_patterns = metadata_patterns or self.METADATA_PATTERNS

    def parse(self, filepath: str) -> ParsedDocument:
        """
        Parse a document (PDF or DOCX).

        Args:
            filepath: Path to the document

        Returns:
            ParsedDocument with extracted text and metadata

        Raises:
            ValueError: If file type is not supported
            FileNotFoundError: If file does not exist
        """
        path = Path(filepath)
        if not path.exists():
            raise FileNotFoundError(f"File not found: {filepath}")

        suffix = path.suffix.lower()
        if suffix == ".pdf":
            return self._parse_pdf(filepath)
        elif suffix == ".docx":
            return self._parse_docx(filepath)
        else:
            raise ValueError(f"Unsupported file type: {suffix}")

    def _parse_pdf(self, filepath: str) -> ParsedDocument:
        """
        Parse a PDF file using pdfplumber.

        Args:
            filepath: Path to the PDF file

        Returns:
            ParsedDocument
        """
        try:
            import pdfplumber
        except ImportError:
            raise ImportError(
                "pdfplumber is required for PDF parsing. "
                "Install with: pip install pdfplumber"
            )

        logger.info(f"Parsing PDF: {filepath}")
        pages: List[ParsedPage] = []

        with pdfplumber.open(filepath) as pdf:
            for i, page in enumerate(pdf.pages):
                # Extract text
                text = page.extract_text() or ""

                # Extract tables if enabled
                tables = []
                if self.extract_tables:
                    raw_tables = page.extract_tables()
                    if raw_tables:
                        tables = [
                            [[cell or "" for cell in row] for row in table]
                            for table in raw_tables
                        ]

                pages.append(ParsedPage(
                    page_number=i + 1,
                    text=text,
                    tables=tables,
                    metadata={},
                ))

        # Create parsed document
        filename = Path(filepath).name
        full_text = "\n\n".join(page.text for page in pages if page.text)

        parsed = ParsedDocument(
            filename=filename,
            file_type="pdf",
            total_pages=len(pages),
            pages=pages,
            full_text=full_text,
        )

        # Extract metadata from first few pages
        self._extract_metadata(parsed)

        logger.info(
            f"Parsed PDF: {len(pages)} pages, {len(full_text)} characters"
        )
        return parsed

    def _parse_docx(self, filepath: str) -> ParsedDocument:
        """
        Parse a DOCX file using python-docx.

        Args:
            filepath: Path to the DOCX file

        Returns:
            ParsedDocument
        """
        try:
            from docx import Document
        except ImportError:
            raise ImportError(
                "python-docx is required for DOCX parsing. "
                "Install with: pip install python-docx"
            )

        logger.info(f"Parsing DOCX: {filepath}")

        doc = Document(filepath)
        pages: List[ParsedPage] = []

        # DOCX doesn't have native page breaks, so we treat it as one page
        # but split on section breaks if present
        paragraphs = []
        current_text = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                current_text.append(text)

            # Check for page break
            if para._element.xml.find("w:lastRenderedPageBreak") != -1:
                if current_text:
                    paragraphs.append("\n".join(current_text))
                    current_text = []

        # Add remaining text
        if current_text:
            paragraphs.append("\n".join(current_text))

        # If no page breaks found, treat as single page
        if not paragraphs:
            paragraphs = ["\n".join(para.text for para in doc.paragraphs if para.text.strip())]

        for i, para_text in enumerate(paragraphs):
            pages.append(ParsedPage(
                page_number=i + 1,
                text=para_text,
                tables=[],
                metadata={},
            ))

        # Extract tables
        if self.extract_tables:
            for table in doc.tables:
                table_data = [
                    [cell.text.strip() for cell in row.cells]
                    for row in table.rows
                ]
                if pages:
                    pages[0].tables.append(table_data)

        # Create parsed document
        filename = Path(filepath).name
        full_text = "\n\n".join(page.text for page in pages if page.text)

        parsed = ParsedDocument(
            filename=filename,
            file_type="docx",
            total_pages=len(pages),
            pages=pages,
            full_text=full_text,
        )

        # Extract metadata
        self._extract_metadata(parsed)

        logger.info(
            f"Parsed DOCX: {len(pages)} pages, {len(full_text)} characters"
        )
        return parsed

    def _extract_metadata(self, doc: ParsedDocument) -> None:
        """
        Extract metadata from document content.

        Looks for tema number, title, and other structured information.

        Args:
            doc: ParsedDocument to extract metadata from
        """
        # Use first few pages for metadata extraction
        text_to_search = "\n".join(
            page.text for page in doc.pages[:3] if page.text
        )

        # Extract tema number
        tema_match = re.search(self.metadata_patterns["tema"], text_to_search)
        if tema_match:
            doc.tema = int(tema_match.group(1))

        # Extract title
        titulo_match = re.search(self.metadata_patterns["titulo"], text_to_search)
        if titulo_match:
            doc.title = titulo_match.group(1).strip()

        # Store raw metadata
        doc.metadata["has_tema"] = doc.tema is not None
        doc.metadata["has_title"] = doc.title is not None

    def extract_text_by_page(self, filepath: str) -> List[Tuple[int, str]]:
        """
        Extract text from each page of a document.

        Args:
            filepath: Path to the document

        Returns:
            List of (page_number, text) tuples
        """
        parsed = self.parse(filepath)
        return [(page.page_number, page.text) for page in parsed.pages]

    def get_text_for_chunking(self, parsed: ParsedDocument) -> List[Dict[str, Any]]:
        """
        Get text segments ready for chunking.

        Args:
            parsed: ParsedDocument

        Returns:
            List of dicts with page_number and text
        """
        segments = []
        for page in parsed.pages:
            if page.text.strip():
                segments.append({
                    "page_number": page.page_number,
                    "text": page.text,
                    "tema": parsed.tema,
                })

                # Add table text as separate segments
                for table in page.tables:
                    table_text = self._table_to_text(table)
                    if table_text.strip():
                        segments.append({
                            "page_number": page.page_number,
                            "text": table_text,
                            "tema": parsed.tema,
                            "is_table": True,
                        })

        return segments

    def _table_to_text(self, table: List[List[str]]) -> str:
        """
        Convert a table to text representation.

        Args:
            table: 2D list of table cells

        Returns:
            Text representation of the table
        """
        if not table:
            return ""

        # Simple text representation
        rows = []
        for row in table:
            rows.append(" | ".join(cell for cell in row if cell))

        return "\n".join(rows)
